"""
Document processing utilities for extracting content from PDF and DOCX files.
"""
import io
import logging
from typing import Tuple, Optional
from fastapi import UploadFile, HTTPException
from PyPDF2 import PdfReader

# Import python-docx with fallback
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    DOCX_AVAILABLE = False

from src.domain.models.llm_selection import ProcessedFileContent
from src.utils.constants import ALLOWED_DOCUMENT_EXTENSIONS

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Utility class for processing PDF and DOCX files.
    """
    
    # Maximum file size: 10MB
    MAX_FILE_SIZE = 10 * 1024 * 1024
    
    @staticmethod
    def validate_file(file: UploadFile) -> None:
        """
        Validate uploaded file.
        
        Args:
            file: Uploaded file to validate
            
        Raises:
            HTTPException: If file validation fails
        """
        # Check if file is provided
        if not file or not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        # Check file extension
        filename = file.filename or ""
        file_extension = filename.lower().split('.')[-1] if '.' in filename else ''
        if f".{file_extension}" not in ALLOWED_DOCUMENT_EXTENSIONS:
            raise HTTPException(
                status_code=400, 
                detail=f"File type not supported. Allowed types: {', '.join(ALLOWED_DOCUMENT_EXTENSIONS)}"
            )
        
        # Check DOCX availability
        if file_extension in ['docx', 'doc'] and not DOCX_AVAILABLE:
            raise HTTPException(
                status_code=500,
                detail="DOCX processing not available. Please install python-docx package."
            )
        
        # Check file size (this is a rough check, actual size check happens during read)
        if hasattr(file, 'size') and file.size and file.size > DocumentProcessor.MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400, 
                detail=f"File too large. Maximum size: {DocumentProcessor.MAX_FILE_SIZE // (1024*1024)}MB"
            )
    
    @staticmethod
    async def extract_pdf_content(file: UploadFile, file_content: bytes) -> ProcessedFileContent:
        """
        Extract text content from PDF file.
        
        Args:
            file: Uploaded PDF file
            file_content: File content as bytes
            
        Returns:
            ProcessedFileContent with extracted text and metadata
            
        Raises:
            HTTPException: If PDF processing fails
        """
        try:
            filename = file.filename or "unknown.pdf"
            
            # Create PDF reader
            file_stream = io.BytesIO(file_content)
            pdf_reader = PdfReader(file_stream)
            
            # Extract text from all pages
            extracted_text = ""
            page_count = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        extracted_text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                except Exception as page_error:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {page_error}")
                    continue
            
            # Clean up extracted text
            extracted_text = extracted_text.strip()
            
            if not extracted_text:
                raise HTTPException(
                    status_code=400, 
                    detail="No text content could be extracted from the PDF file"
                )
            
            logger.info(f"Successfully extracted {len(extracted_text)} characters from PDF: {filename}")
            
            return ProcessedFileContent(
                filename=filename,
                content=extracted_text,
                file_type="pdf",
                file_size=len(file_content),
                page_count=page_count
            )
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Error processing PDF file {file.filename}: {e}")
            raise HTTPException(
                status_code=500, 
                detail=f"Failed to process PDF file: {str(e)}"
            )

    @staticmethod
    async def extract_docx_content(file: UploadFile, file_content: bytes) -> ProcessedFileContent:
        """
        Extract text content from DOCX file.
        
        Args:
            file: Uploaded DOCX file
            file_content: File content as bytes
            
        Returns:
            ProcessedFileContent with extracted text and metadata
            
        Raises:
            HTTPException: If DOCX processing fails
        """
        try:
            if not DOCX_AVAILABLE:
                raise HTTPException(
                    status_code=500,
                    detail="DOCX processing not available. Please install python-docx package."
                )
            
            filename = file.filename or "unknown.docx"
            
            # Create DOCX document
            file_stream = io.BytesIO(file_content)
            doc = Document(file_stream)
            
            # Extract text from all paragraphs
            extracted_text = ""
            paragraph_count = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    extracted_text += paragraph.text + "\n"
                    paragraph_count += 1
            
            # Extract text from tables if any
            table_count = 0
            for table in doc.tables:
                table_count += 1
                extracted_text += f"\n--- Table {table_count} ---\n"
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        extracted_text += " | ".join(row_text) + "\n"
            
            # Clean up extracted text
            extracted_text = extracted_text.strip()
            
            if not extracted_text:
                raise HTTPException(
                    status_code=400, 
                    detail="No text content could be extracted from the DOCX file"
                )
            
            logger.info(f"Successfully extracted {len(extracted_text)} characters from DOCX: {filename}")
            
            return ProcessedFileContent(
                filename=filename,
                content=extracted_text,
                file_type="docx",
                file_size=len(file_content),
                page_count=paragraph_count  # Use paragraph count as a page equivalent
            )
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Error processing DOCX file {file.filename}: {e}")
            raise HTTPException(
                status_code=500, 
                detail=f"Failed to process DOCX file: {str(e)}"
            )
    
    @staticmethod
    async def extract_document_content(file: UploadFile) -> ProcessedFileContent:
        """
        Extract text content from uploaded document (PDF or DOCX).
        
        Args:
            file: Uploaded document file
            
        Returns:
            ProcessedFileContent with extracted text and metadata
            
        Raises:
            HTTPException: If document processing fails
        """
        try:
            # Validate file first
            DocumentProcessor.validate_file(file)
            
            # Read file content
            file_content = await file.read()
            file_size = len(file_content)
            
            # Check actual file size
            if file_size > DocumentProcessor.MAX_FILE_SIZE:
                raise HTTPException(
                    status_code=400, 
                    detail=f"File too large. Maximum size: {DocumentProcessor.MAX_FILE_SIZE // (1024*1024)}MB"
                )
            
            # Determine file type and process accordingly
            filename = file.filename or ""
            file_extension = filename.lower().split('.')[-1] if '.' in filename else ''
            
            if file_extension == 'pdf':
                return await DocumentProcessor.extract_pdf_content(file, file_content)
            elif file_extension in ['docx', 'doc']:
                return await DocumentProcessor.extract_docx_content(file, file_content)
            else:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file type: {file_extension}"
                )
            
        except HTTPException:
            # Re-raise HTTP exceptions
            raise
        except Exception as e:
            logger.error(f"Error processing document file {file.filename}: {e}")
            raise HTTPException(
                status_code=500, 
                detail=f"Failed to process document file: {str(e)}"
            )
        finally:
            # Reset file pointer for potential future use
            if hasattr(file, 'seek'):
                try:
                    await file.seek(0)
                except:
                    pass

    @staticmethod
    def generate_file_insights_prompt(file_content: ProcessedFileContent, user_query: Optional[str] = None) -> str:
        """
        Generate appropriate prompt for LLM based on file content and optional query.
        
        Args:
            file_content: Processed file content
            user_query: Optional user query
            
        Returns:
            Formatted prompt for LLM
        """
        file_type_display = file_content.file_type.upper()
        if user_query and user_query.strip():
            # User provided both file and query
            prompt = f"""I have uploaded a {file_type_display} file named "{file_content.filename}" ({file_content.page_count} pages/sections, {file_content.file_size} bytes) and have a specific question about it.

File Content:
{file_content.content}

My Question: {user_query.strip()}

Please answer my question based on the document content above."""

        else:
            # Auto-generate comprehensive insights
            prompt = f"""I have uploaded a {file_type_display} document named "{file_content.filename}" ({file_content.page_count} pages/sections, {file_content.file_size} bytes). Please provide a comprehensive analysis including:

1. Document Summary - Brief overview of the main content
2. Key Topics & Themes - Main subjects and recurring themes
3. Important Information - Critical data, findings, or conclusions
4. Structure & Organization - How the document is organized
5. Key Takeaways - Most important points to remember
6. Actionable Items - Any tasks, recommendations, or next steps mentioned
7. Context & Significance - Why this document matters and its broader implications

Document Content:
{file_content.content}

Please provide a thorough analysis covering all the points above."""

        return prompt

    @staticmethod
    def create_file_summary_for_history(file_content: ProcessedFileContent) -> str:
        """
        Create a brief summary of the file for conversation history.
        
        Args:
            file_content: Processed file content
            
        Returns:
            Brief summary string
        """
        content_preview = file_content.content[:200] + "..." if len(file_content.content) > 200 else file_content.content
        return f"[{file_content.file_type.upper()} File: {file_content.filename}, {file_content.page_count} pages/sections] Preview: {content_preview}"


# Maintain backward compatibility
PDFProcessor = DocumentProcessor 